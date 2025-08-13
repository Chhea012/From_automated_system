from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io

def generate_docx(data):
    # Create document
    doc = Document()
    
    # Set default font and style
    style = doc.styles['Normal']
    style.font.name = 'Calibri (Body)'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(5)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style.font.color.rgb = RGBColor(0, 0, 0)

    # Define dynamic bold variables for Party A and Party B
    def add_bold_party(paragraph, text):
        run = paragraph.add_run(text)
        run.bold = True
        return run

    # Title Section
    p = doc.add_paragraph('The Service Agreement')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = 'Calibri (Body)'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph('ON')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph(data['project_title'])
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = 'Calibri (Body)'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph(f"No.: {data['contract_number']}")
    p.runs[0].bold = True
    p.runs[0].font.name = 'Calibri (Body)'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(14)
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph('BETWEEN')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)

    # Party A
    p = doc.add_paragraph()
    run = p.add_run(f"{data['organization_name']}, represented by ")
    run.bold = True
    run = p.add_run(f"{data['party_a_name']}, ")
    run.bold = True
    run = p.add_run(f"{data['party_a_position']}.\n")
    run.bold = False
    run = p.add_run(f"Address: {data['party_a_address']}.\n")
    run.bold = False
    run = p.add_run("hereinafter called the “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("”")
    run.bold = False
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph('AND')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)

    # Party B with email in blue
    p = doc.add_paragraph()
    run = p.add_run(f"{data['party_b_full_name_with_title']},\n")
    run.bold = True
    run = p.add_run(f"Address: {data['party_b_address']}\n")
    run.bold = False
    run = p.add_run(f"H/P: {data['party_b_phone']}, E-mail: ")
    run.bold = False
    email_run = p.add_run(data['party_b_email'])
    email_run.font.color.rgb = RGBColor(0, 0, 255)
    email_run.bold = False
    run = p.add_run("\nhereinafter called the “")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("”")
    run.bold = False
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Whereas Clauses
    p = doc.add_paragraph(
        f"Whereas {data['organization_name']} is a legal entity registered with the Ministry of Interior (MOI) {data['registration_number']} dated {data['registration_date']}."
    )
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("Whereas NGOF will engage the services of “")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” which accept the engagement under the following term and conditions.")
    run.bold = False
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph("Both Parties Agreed as follows:")
    p.runs[0].bold = True
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Article 1
    p = doc.add_heading('', level=2)
    run = p.add_run('ARTICLE 1')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p.add_run(': TERMS OF REFERENCE')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    p = doc.add_paragraph()
    run = p.add_run("“")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” shall perform tasks as stated in the attached TOR (annex-1) to “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("”, and deliver each milestone as stipulated in article 4.\n")
    run.bold = False
    run = p.add_run("The work shall be of good quality and well performed with the acceptance by “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run(".”")
    run.bold = False

    # Article 2
    p = doc.add_heading('', level=2)
    run = p.add_run('ARTICLE 2')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p.add_run(': TERM OF AGREEMENT')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    p = doc.add_paragraph()
    run = p.add_run(f"The agreement is effective from {data['agreement_start_date']} – {data['agreement_end_date']}.")
    run.bold = True
    run = p.add_run(
        " This Agreement is terminated automatically after the due date of the Agreement Term unless otherwise, "
        "both Parties agree to extend the Term with a written agreement."
    )
    run.bold = False

    # Article 3
    p = doc.add_heading('', level=2)
    run = p.add_run('ARTICLE 3')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p.add_run(': PROFESSIONAL FEE')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    try:
        total_fee = float(data['total_fee_usd'])
        tax_percentage = float(data['tax_percentage'])
    except (ValueError, TypeError) as e:
        raise ValueError(f"Invalid numeric values for total_fee_usd or tax_percentage: {str(e)}")
    total_fee_formatted = f"{total_fee:.2f}"
    tax_amount = total_fee * (tax_percentage / 100)
    net_amount = total_fee - tax_amount
    p = doc.add_paragraph()
    run = p.add_run(f"The professional fee is the total amount of USD {total_fee_formatted} ({data['total_fee_words']}) including tax for the whole assignment period.")
    run.bold = True
    p = doc.add_paragraph(f"    Total Service Fee:        USD {total_fee_formatted}")
    p.runs[0].bold = True
    p = doc.add_paragraph(f"    Withholding Tax {tax_percentage:.1f}%:    USD {tax_amount:.2f}")
    p.runs[0].bold = True
    p = doc.add_paragraph(f"    Net amount:            USD {net_amount:.2f}")
    p.runs[0].bold = True
    p = doc.add_paragraph()
    run = p.add_run("“")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” is responsible to issue the Invoice (net amount) and receipt (when receiving the payment) with the total amount as stipulated in each instalment as in the Article 4 after having done the agreed deliverable tasks, for payment request. ")
    run.bold = False
    run = p.add_run("The payment will be processed after the satisfaction from “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("” as of the required deliverable tasks as stated in Article 4.")
    run.bold = False
    p = doc.add_paragraph()
    run = p.add_run("“")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” is responsible for all related taxes payable to the government department.")
    run.bold = False

    # Article 4
    p = doc.add_heading('', level=2)
    run = p.add_run('ARTICLE 4')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p.add_run(': TERM OF PAYMENT')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    p = doc.add_paragraph('The payment will be made based on the following schedules:')

    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(1.2)
    table.columns[1].width = Inches(1.5)
    table.columns[2].width = Inches(3.5)
    table.columns[3].width = Inches(1.2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Installment'
    hdr_cells[1].text = 'Total Amount (USD)'
    hdr_cells[2].text = 'Deliverable'
    hdr_cells[3].text = 'Due date'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.size = Pt(12)
        cell.paragraphs[0].runs[0].font.name = 'Calibri (Body)'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].bold = True
    row_cells = table.rows[1].cells
    row_cells[0].text = data['payment_installment_desc']
    row_cells[1].text = f"· Gross: ${total_fee_formatted}\n· Tax {tax_percentage:.1f}%: ${tax_amount:.2f}\n· Net pay: ${net_amount:.2f}"
    row_cells[1].paragraphs[0].runs[0].bold = True
    deliverables_list = "\n".join([f"· {line.strip()}" for line in data['deliverables'].split('\n') if line.strip()])
    row_cells[2].text = deliverables_list
    row_cells[3].text = data['agreement_end_date']
    for cell in row_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in cell.paragraphs[0].runs:
            run.font.name = 'Calibri (Body)'
            if cell != row_cells[1]:  # Only the Total Amount cell should be bold
                run.bold = False

    # Article 5
    p = doc.add_heading('', level=2)
    run = p.add_run('ARTICLE 5')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p.add_run(': NO OTHER PERSONS')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    p = doc.add_paragraph(
        "No person or entity, which is not a party to this agreement, has any rights to enforce, take any action, or claim it is owed any benefit under this agreement."
    )

    # Article 6
    p = doc.add_heading('', level=2)
    run = p.add_run('ARTICLE 6')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p.add_run(': MONITORING and COORDINATION')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    p = doc.add_paragraph()
    run = p.add_run("“")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("” shall monitor and evaluate the progress of the agreement toward its objective, including the activities implemented. ")
    run.bold = False
    run = p.add_run(f"{data['focal_person_a_name']}, ")
    run.bold = True
    run = p.add_run(f"{data['focal_person_a_position']} ")
    run.bold = True
    run = p.add_run("(Telephone ")
    run.bold = False
    run = p.add_run(f"{data['focal_person_a_phone']} ")
    run.bold = True
    run = p.add_run("Email: ")
    run.bold = False
    email_run = p.add_run(data['focal_person_a_email'])
    email_run.font.color.rgb = RGBColor(0, 0, 255)
    email_run.bold = False
    run = p.add_run(") is the focal contact person of “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("” and ")
    run.bold = False
    run = p.add_run(f"{data['party_b_full_name_with_title']} ")
    run.bold = True
    run = p.add_run("(HP. ")
    run.bold = False
    run = p.add_run(f"{data['party_b_phone']}, ")
    run.bold = True
    run = p.add_run("E-mail: ")
    run.bold = False
    email_run = p.add_run(data['party_b_email'])
    email_run.font.color.rgb = RGBColor(0, 0, 255)
    email_run.bold = False
    run = p.add_run(") the focal contact person of “")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("”. The focal contact person of “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("” and “")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” will work together for overall coordination including reviewing and meeting discussions during the assignment process.")
    run.bold = False

    # Article 7
    p = doc.add_heading('', level=2)
    run = p.add_run('ARTICLE 7')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p.add_run(': CONFIDENTIALITY')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    p = doc.add_paragraph()
    run = p.add_run(f"All outputs produced, with the exception of the “{data['output_description']}”, which is a contribution from, and to be claimed as a public document by the main author and co-author in associated, and/or under this agreement, shall be the property of “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("”. ")
    run.bold = False
    run = p.add_run("The “")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” agrees to not disclose any confidential information, of which he/she may take cognizance in the performance under this contract, except with the prior written approval of the “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("”.")
    run.bold = False

    # Article 8
    p = doc.add_heading('', level=2)
    run = p.add_run('ARTICLE 8')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p.add_run(': ANTI-CORRUPTION and CONFLICT OF INTEREST')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    p = doc.add_paragraph()
    run = p.add_run("“")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” shall not participate in any practice that is or could be construed as an illegal or corrupt practice in Cambodia. ")
    run.bold = False
    run = p.add_run("The “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("” is committed to fighting all types of corruption and expects this same commitment from the consultant it reserves the rights and believes based on the declaration of “")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” that it is an independent social enterprise firm operating in Cambodia and it does not involve any conflict of interest with other parties that may be affected to the “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("”.")
    run.bold = False

    # Article 9
    p = doc.add_heading('', level=2)
    run = p.add_run('ARTICLE 9')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p.add_run(': OBLIGATION TO COMPLY WITH THE NGOF’S POLICIES AND CODE OF CONDUCT')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    p = doc.add_paragraph()
    run = p.add_run("By signing this agreement, “")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” is obligated to comply with and respect all existing policies and code of conduct of “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("”, such as Gender Mainstreaming, Child Protection, Disability policy, Environmental Mainstreaming, etc. and the “")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” declared themselves that s/he will perform the assignment in the neutral position, professional manner, and not be involved in any political affiliation.")
    run.bold = False

    # Article 10
    p = doc.add_heading('', level=2)
    run = p.add_run('ARTICLE 10')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p.add_run(': ANTI-TERRORISM FINANCING AND FINANCIAL CRIME')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    p = doc.add_paragraph(
        'NGOF is determined that all its funds and resources should only be used to further its mission and shall not be subject to illicit use by any third party nor used or abused for any illicit purpose. '
        'In order to achieve this objective, NGOF will not knowingly or recklessly provide funds, economic goods, or material support to any entity or individual designated as a “terrorist” by the international community or affiliate domestic governments and will take all reasonable steps to safeguard and protect its assets from such illicit use and to comply with host government laws.\n'
        'NGOF respects its contracts with its donors and puts procedures in place for compliance with these contracts.\n'
        '“Illicit use” refers to terrorist financing, sanctions, money laundering, and export control regulations.'
    )

    # Article 11
    p = doc.add_heading('', level=2)
    run = p.add_run('ARTICLE 11')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p.add_run(': INSURANCE')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    p = doc.add_paragraph()
    run = p.add_run("“")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” is responsible for any health and life insurance of its team members. “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("” will not be held responsible for any medical expenses or compensation incurred during or after this contract.")
    run.bold = False

    # Article 12
    p = doc.add_heading('', level=2)
    run = p.add_run('ARTICLE 12')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p.add_run(': ASSIGNMENT')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    p = doc.add_paragraph()
    run = p.add_run("“")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” shall have the right to assign individuals within its organization to carry out the tasks herein named in the attached Technical Proposal. ")
    run.bold = False
    run = p.add_run("The “")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” shall not assign, or transfer any of its rights or obligations under this agreement hereunder without the prior written consent of “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("”. ")
    run.bold = False
    run = p.add_run("Any attempt by “")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” to assign or transfer any of its rights and obligations without the prior written consent of “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("” shall render this agreement subject to immediate termination by “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("”.")
    run.bold = False

    # Article 13
    p = doc.add_heading('', level=2)
    run = p.add_run('ARTICLE 13')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p.add_run(': RESOLUTION OF CONFLICTS/DISPUTES')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    p = doc.add_paragraph()
    run = p.add_run("Conflicts between any of these agreements shall be resolved by the following methods:\n")
    run.bold = False
    run = p.add_run("In the case of a disagreement arising between “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("” and the “")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” regarding the implementation of any part of, or any other substantive question arising under or relating to this agreement, the parties shall use their best efforts to arrive at an agreeable resolution by mutual consultation.\n")
    run.bold = False
    run = p.add_run("Unresolved issues may, upon the option of either party and written notice to the other party, be referred to for arbitration. ")
    run.bold = False
    run = p.add_run("Failure by the “")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” or “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("” to dispute a decision arising from such arbitration in writing within thirty (30) calendar days of receipt of a final decision shall result in such final decision being deemed binding upon either the “")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” and/or “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("”. ")
    run.bold = False
    run = p.add_run("All expenses related to arbitration will be shared equally between both parties.")
    run.bold = False

    # Article 14
    p = doc.add_heading('', level=2)
    run = p.add_run('ARTICLE 14')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p.add_run(': TERMINATION')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    p = doc.add_paragraph()
    run = p.add_run("The “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("” or the “")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” may, by notice in writing, terminate this agreement under the following conditions:\n")
    run.bold = False
    run = p.add_run("1. “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("” may terminate this agreement at any time with a week notice if “")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” fails to comply with the terms and conditions of this agreement.\n")
    run.bold = False
    run = p.add_run("2. For gross professional misconduct (as defined in the NGOF Human Resource Policy), “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("” may terminate this agreement immediately without prior notice. ")
    run.bold = False
    run = p.add_run("“")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("” will notify “")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” in a letter that will indicate the reason for termination as well as the effective date of termination.\n")
    run.bold = False
    run = p.add_run("3. “")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” may terminate this agreement at any time with a one-week notice if “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("” fails to comply with the terms and conditions of this agreement. ")
    run.bold = False
    run = p.add_run("“")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” will notify “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("” in a letter that will indicate the reason for termination as well as the effective date of termination. ")
    run.bold = False
    run = p.add_run("But if “")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” intended to terminate this agreement by itself without any appropriate reason or fails of implementing the assignment, “")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” has to refund the full amount of fees received to “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("”.\n")
    run.bold = False
    run = p.add_run("4. If for any reason either “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("” or the “")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” decides to terminate this agreement, “")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” shall be paid pro-rata for the work already completed by “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("”. ")
    run.bold = False
    run = p.add_run("This payment will require the submission of a timesheet that demonstrates work completed as well as the handing over of any deliverables completed or partially completed. ")
    run.bold = False
    run = p.add_run("In case “")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” has received payment for services under the agreement which have not yet been performed; the appropriate portion of these fees would be refunded by “")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("” to “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("”.")
    run.bold = False

    # Article 15
    p = doc.add_heading('', level=2)
    run = p.add_run('ARTICLE 15')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p.add_run(': MODIFICATION OR AMENDMENT')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    p = doc.add_paragraph()
    run = p.add_run("No modification or amendment of this agreement shall be valid unless in writing and signed by an authorized person of “")
    run.bold = False
    add_bold_party(p, "Party A")
    run = p.add_run("” and “")
    run.bold = False
    add_bold_party(p, "Party B")
    run = p.add_run("”.")
    run.bold = False

    # Article 16
    p = doc.add_heading('', level=2)
    run = p.add_run('ARTICLE 16')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p.add_run(': CONTROLLING OF LAW')
    run.font.size = Pt(11)
    run.font.name = 'Calibri (Body)'
    run.bold = True
    run.underline = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    p = doc.add_paragraph(
        'This agreement shall be governed and construed following the law of the Kingdom of Cambodia. '
        'The Simultaneous Interpretation Agreement is prepared in two original copies.'
    )

    # Date
    p = doc.add_paragraph(f"Date: {data['agreement_start_date']}")
    p.runs[0].bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Signature Table
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(3.0)
    table.columns[1].width = Inches(3.0)
    col1, col2 = table.rows[0].cells
    col1.paragraphs[0].add_run('For “Party A”\n\n\n_________________\n').bold = True
    col1.paragraphs[0].add_run(data['party_a_signature_name'] + '\n').bold = True
    col1.paragraphs[0].add_run(data['party_a_position']).bold = True
    col1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    col2.paragraphs[0].add_run('For “Party B”\n\n\n____________________\n').bold = True
    col2.paragraphs[0].add_run(data['party_b_signature_name'] + '\n').bold = True
    col2.paragraphs[0].add_run(data['party_b_position']).bold = True
    col2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add page break to ensure multi-page layout
    doc.add_page_break()

    # Save to BytesIO for portability
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    
    return docx_bytes